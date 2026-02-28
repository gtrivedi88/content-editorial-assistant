"""DOCX parser using python-docx for paragraph, table, and style extraction.

Maps Word paragraph styles to semantic block types and extracts tables,
lists, and images from ``.docx`` files.
"""

import logging
import zipfile
from typing import Optional

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph

from app.services.parsing.base import BaseParser, Block, ParseResult

logger = logging.getLogger(__name__)

# Style name prefix -> block_type mapping
_HEADING_PREFIX = "Heading"
_CODE_STYLES = frozenset({"Code", "CodeBlock", "Source Code", "code", "Verbatim"})
_LIST_STYLES = frozenset({
    "List Bullet", "List Number", "List Paragraph",
    "List Bullet 2", "List Bullet 3",
    "List Number 2", "List Number 3",
})


class DocxParser(BaseParser):
    """Parser for Microsoft Word ``.docx`` documents.

    Uses python-docx to iterate paragraphs and tables, mapping paragraph
    styles to semantic block types.
    """

    def parse(self, content: str, filename: Optional[str] = None) -> ParseResult:
        """Parse a ``.docx`` file and return structured blocks.

        *content* is treated as a file path to the ``.docx`` file.

        Args:
            content: File path to the ``.docx`` document.
            filename: Optional display name for metadata.

        Returns:
            ParseResult with text blocks extracted from the document.
        """
        logger.debug("DocxParser: opening %s", filename or content)

        if not content:
            return ParseResult(blocks=[], plain_text="")

        try:
            doc = Document(content)
        except PackageNotFoundError:
            logger.warning("DocxParser: file not found: %s", content)
            return ParseResult(
                blocks=[], plain_text="",
                metadata={"error": "File not found"},
            )
        except KeyError as exc:
            logger.warning("DocxParser: missing expected part in docx: %s", exc)
            return ParseResult(
                blocks=[], plain_text="",
                metadata={"error": f"Invalid docx structure: {exc}"},
            )
        except ValueError as exc:
            logger.warning("DocxParser: value error opening docx: %s", exc)
            return ParseResult(
                blocks=[], plain_text="",
                metadata={"error": str(exc)},
            )
        except zipfile.BadZipFile:
            logger.warning("DocxParser: not a valid ZIP / docx file: %s", content)
            return ParseResult(
                blocks=[], plain_text="",
                metadata={"error": "Not a valid docx file"},
            )

        blocks = self._extract_blocks(doc)

        plain_text = "\n\n".join(
            b.content for b in blocks if not b.should_skip_analysis and b.content
        )

        logger.debug("DocxParser: produced %d blocks", len(blocks))
        return ParseResult(
            blocks=blocks,
            plain_text=plain_text,
            metadata={"filename": filename},
        )

    # ------------------------------------------------------------------
    # Block extraction
    # ------------------------------------------------------------------

    def _extract_blocks(self, doc: Document) -> list[Block]:
        """Walk document body elements and emit blocks in order."""
        blocks: list[Block] = []
        offset = 0

        for element in doc.element.body:
            tag = _local_tag(element)

            if tag == "p":
                para = _find_paragraph(doc, element)
                if para is not None:
                    block = self._paragraph_to_block(para, offset)
                    if block is not None:
                        blocks.append(block)
                        offset = block.end_pos + 1

            elif tag == "tbl":
                table = _find_table(doc, element)
                if table is not None:
                    block = self._table_to_block(table, offset)
                    if block is not None:
                        blocks.append(block)
                        offset = block.end_pos + 1

        return blocks

    @staticmethod
    def _paragraph_to_block(para: Paragraph, offset: int) -> Optional[Block]:
        """Convert a python-docx Paragraph to a Block."""
        text = para.text.strip()

        # Detect images (runs containing inline shapes)
        has_image = _paragraph_has_image(para)
        if has_image and not text:
            end = offset + 10  # placeholder length
            return Block(
                block_type="image",
                content="[image]",
                raw_content="[image]",
                start_pos=offset,
                end_pos=end,
                metadata={"style": _style_name(para)},
            )

        if not text:
            return None

        style_name = _style_name(para)
        block_type, level = _classify_style(style_name)
        skip = block_type == "code_block"
        end = offset + len(text)

        return Block(
            block_type=block_type,
            content=text,
            raw_content=text,
            start_pos=offset,
            end_pos=end,
            level=level,
            should_skip_analysis=skip,
            metadata={"style": style_name},
        )

    @staticmethod
    def _table_to_block(table: DocxTable, offset: int) -> Optional[Block]:
        """Convert a python-docx Table to a table Block with row/cell children."""
        rows: list[Block] = []
        row_offset = offset
        all_text_parts: list[str] = []

        for row in table.rows:
            cells: list[Block] = []
            cell_offset = row_offset
            row_text_parts: list[str] = []

            for cell in row.cells:
                cell_text = cell.text.strip()
                row_text_parts.append(cell_text)
                cell_end = cell_offset + max(len(cell_text), 1)
                cells.append(Block(
                    block_type="table_cell",
                    content=cell_text,
                    raw_content=cell_text,
                    start_pos=cell_offset,
                    end_pos=cell_end,
                ))
                cell_offset = cell_end + 1

            row_text = " | ".join(row_text_parts)
            all_text_parts.append(row_text)
            row_end = cell_offset
            rows.append(Block(
                block_type="table_row",
                content=row_text,
                raw_content=row_text,
                start_pos=row_offset,
                end_pos=row_end,
                children=cells,
            ))
            row_offset = row_end + 1

        if not rows:
            return None

        content = "\n".join(all_text_parts)
        end = row_offset
        return Block(
            block_type="table",
            content=content,
            raw_content=content,
            start_pos=offset,
            end_pos=end,
            children=rows,
        )


# ------------------------------------------------------------------
# Pure helper functions
# ------------------------------------------------------------------


def _local_tag(element: object) -> str:
    """Return the local tag name of an lxml element (stripping namespace)."""
    tag = getattr(element, "tag", "")
    if not isinstance(tag, str):
        return ""
    if "}" in tag:
        return tag.split("}", 1)[1]
    return tag


def _style_name(para: Paragraph) -> str:
    """Return the style name of a paragraph, or empty string."""
    if para.style and para.style.name:
        return para.style.name
    return ""


def _classify_style(style_name: str) -> tuple[str, int]:
    """Map a Word style name to (block_type, level).

    Returns:
        Tuple of block_type string and heading level (0 for non-headings).
    """
    if style_name.startswith(_HEADING_PREFIX):
        suffix = style_name[len(_HEADING_PREFIX):].strip()
        try:
            level = int(suffix)
        except ValueError:
            level = 1
        return "heading", level

    if style_name in _CODE_STYLES:
        return "code_block", 0

    if style_name in _LIST_STYLES:
        return "list_item", 0

    return "paragraph", 0


def _paragraph_has_image(para: Paragraph) -> bool:
    """Return True if any run in *para* contains an inline image."""
    for run in para.runs:
        if run.element is not None:
            drawings = run.element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
            )
            if not drawings:
                drawings = run.element.findall(
                    ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                )
            if drawings:
                return True
    return False


def _find_paragraph(doc: Document, element: object) -> Optional[Paragraph]:
    """Find the python-docx Paragraph that wraps *element*."""
    for para in doc.paragraphs:
        if para._element is element:
            return para
    return None


def _find_table(doc: Document, element: object) -> Optional[DocxTable]:
    """Find the python-docx Table that wraps *element*."""
    for table in doc.tables:
        if table._element is element:
            return table
    return None


